# -*- coding: utf-8 -*-
"""
اختبار ميزة طباعة Word
"""

def test_word_creation():
    """اختبار إنشاء ملف Word"""
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        print("جميع المكتبات المطلوبة متوفرة")
        
        # إنشاء مستند تجريبي
        doc = Document()
        
        # إضافة عنوان
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('اختبار قائمة الطلاب')
        title_run.font.size = Pt(18)
        title_run.bold = True
        
        # إضافة جدول تجريبي
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # إضافة عناوين
        header_row = table.rows[0]
        header_row.cells[0].text = 'المعرف'
        header_row.cells[1].text = 'الاسم'
        header_row.cells[2].text = 'الصف'
        
        # إضافة بيانات تجريبية
        row = table.add_row()
        row.cells[0].text = '1'
        row.cells[1].text = 'أحمد محمد'
        row.cells[2].text = 'الرابع الابتدائي'
        
        # حفظ الملف
        import tempfile
        import os
        from datetime import datetime
        
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"اختبار_قائمة_الطلاب_{timestamp}.docx"
        filepath = os.path.join(temp_dir, filename)
        
        doc.save(filepath)
        
        print(f"تم إنشاء ملف Word بنجاح: {filepath}")
        
        # محاولة فتح الملف
        if os.path.exists(filepath):
            os.startfile(filepath)
            print("تم فتح الملف في Word")
        
        return True
        
    except Exception as e:
        print(f"خطأ في اختبار Word: {e}")
        import traceback
        print(traceback.format_exc())
        return False


if __name__ == "__main__":
    print("بدء اختبار إنشاء ملف Word...")
    test_word_creation()
