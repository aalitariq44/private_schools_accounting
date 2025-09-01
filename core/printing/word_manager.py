# -*- coding: utf-8 -*-
"""
مدير إنشاء ملفات Word للطباعة
"""

import logging
import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
import subprocess

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("مكتبة python-docx غير متوفرة، لن تتمكن من إنشاء ملفات Word")

from PyQt5.QtWidgets import QMessageBox, QFileDialog
import config


class WordManager:
    """مدير إنشاء ملفات Word"""
    
    def __init__(self, parent=None):
        self.parent = parent
        
        if not DOCX_AVAILABLE:
            raise ImportError("مكتبة python-docx غير متوفرة")
    
    def create_students_list_word(self, students: List[Dict], selected_columns: Dict, filter_info: Optional[str] = None):
        """إنشاء قائمة الطلاب في ملف Word"""
        try:
            # إنشاء مستند Word جديد
            doc = Document()
            # اجعل اتجاه الكتابة من اليمين لليسار بشكل افتراضي
            from docx.oxml import OxmlElement
            normal_style = doc.styles['Normal']
            # ضبط فقرة RTL
            pPr = normal_style.element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            # ضبط النص RTL
            rPr = normal_style.element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rPr.append(rtl)
            
            # إعداد اتجاه الصفحة واللغة
            section = doc.sections[0]
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            # ضبط اتجاه النص من اليمين لليسار
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(11)
            
            # إضافة عنوان رئيسي
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run('قائمة الطلاب')
            title_run.font.size = Pt(18)
            title_run.bold = True
            title_run.font.name = 'Arial'
            
            # إضافة تاريخ الطباعة
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_paragraph.add_run(f'تاريخ الطباعة: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            date_run.font.size = Pt(12)
            date_run.font.name = 'Arial'
            
            # إضافة فراغ
            doc.add_paragraph()
            
            # عرض الفلاتر المطبقة وإجمالي العدد في جدول صغير بمحاذاة اليمين
            if filter_info:
                info_tbl = doc.add_table(rows=2, cols=1)
            else:
                info_tbl = doc.add_table(rows=1, cols=1)
            info_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            # الصف الأول: الفلاتر
            if filter_info:
                cell0 = info_tbl.cell(0, 0)
                p0 = cell0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run0 = p0.add_run(f'الفلاتر المطبقة: {filter_info}')
                run0.font.size = Pt(10)
                run0.italic = True
                run0.font.name = 'Arial'
            # الصف الثاني أو الوحيد: الملخص
            idx = 1 if filter_info else 0
            cell1 = info_tbl.cell(idx, 0)
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1 = p1.add_run(f'إجمالي عدد الطلاب: {len(students)}')
            run1.font.size = Pt(12)
            run1.bold = True
            run1.font.name = 'Arial'
            # مسافة بعد الجدول الصغير
            doc.add_paragraph()
            
            # إنشاء الجدول
            if students:
                # تحديد عدد الأعمدة (ت + الأعمدة المحددة)
                num_cols = len(selected_columns) + 1
                table = doc.add_table(rows=1, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                # جعل صف العنوان يتكرر في كل صفحة
                hdr_row = table.rows[0]
                # إضافة وسم tblHeader لعنصر الصف
                hdr_tr = hdr_row._tr
                trPr = hdr_tr.get_or_add_trPr()
                tblHeader = OxmlElement('w:tblHeader')
                trPr.append(tblHeader)
                
                # إضافة عناوين الأعمدة (RTL: العنوان الأخير على اليمين)
                header_row = table.rows[0]
                # ضع عمود الترقيم في الخلية الأخيرة
                header_row.cells[num_cols-1].text = 'ت'
                
                # ترتيب الأعمدة كما في المتطلبات
                column_order = ['id', 'name', 'school_name', 'grade', 'section', 'gender', 'phone', 'status', 'total_fee', 'total_paid', 'remaining']
                ordered_columns = []
                
                for col_key in column_order:
                    if col_key in selected_columns:
                        ordered_columns.append((col_key, selected_columns[col_key]))
                
                # إضافة أي أعمدة أخرى لم تكن في الترتيب الأساسي
                for col_key, col_name in selected_columns.items():
                    if col_key not in [item[0] for item in ordered_columns]:
                        ordered_columns.append((col_key, col_name))
                
                # ملء عناوين الأعمدة بالعكس (من اليمين لليسار)
                for i, (col_key, col_name) in enumerate(ordered_columns):
                    # تبدأ الأعمدة المحددة قبل عمود الترقيم، من الخلية قبل الأخيرة إلى الأولى
                    header_row.cells[num_cols - 2 - i].text = col_name
                
                # تنسيق عناوين الأعمدة
                for cell in header_row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # التأكد من وجود run قبل الوصول إليه
                    if cell.paragraphs[0].runs:
                        run = cell.paragraphs[0].runs[0]
                    else:
                        run = cell.paragraphs[0].add_run()
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
                
                # إضافة بيانات الطلاب (RTL: الترقيم في اليمين وبقية الأعمدة تتجه نحو اليسار)
                for index, student in enumerate(students, 1):
                    row = table.add_row()
                    # ضع الترقيم في الخلية الأخيرة
                    row.cells[num_cols-1].text = str(index)
                    for i, (col_key, col_name) in enumerate(ordered_columns):
                        cell_value = student.get(col_key, '')
                        if cell_value is None:
                            cell_value = ''
                        # املأ الأعمدة المحددة بالعكس من قبل عمود الترقيم
                        row.cells[num_cols - 2 - i].text = str(cell_value)
                    
                    # تنسيق صف البيانات
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # التأكد من وجود runs قبل الوصول إليها
                        if cell.paragraphs[0].runs:
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Arial'
                
                # ضبط عرض الأعمدة
                for column in table.columns:
                    for cell in column.cells:
                        cell.width = Inches(1.2)
            
            else:
                # في حالة عدم وجود طلاب
                no_data_paragraph = doc.add_paragraph()
                no_data_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                no_data_run = no_data_paragraph.add_run('لا توجد بيانات طلاب للعرض')
                no_data_run.font.size = Pt(14)
                no_data_run.font.name = 'Arial'
            
            # إضافة تذييل
            doc.add_paragraph()
            footer_paragraph = doc.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_paragraph.add_run('نظام غريديا لادارة حسابات المدارس الاهلية')
            footer_run.font.size = Pt(10)
            footer_run.italic = True
            footer_run.font.name = 'Arial'
            
            # حفظ الملف
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"قائمة_الطلاب_{timestamp}.docx"
            filepath = os.path.join(temp_dir, filename)
            
            doc.save(filepath)
            
            # فتح الملف
            self.open_word_file(filepath)
            
            logging.info(f"تم إنشاء ملف Word بنجاح: {filepath}")
            return filepath
            
        except Exception as e:
            logging.error(f"خطأ في إنشاء ملف Word: {e}")
            if self.parent:
                QMessageBox.critical(self.parent, "خطأ", f"فشل في إنشاء ملف Word:\n{str(e)}")
            return None
    
    def open_word_file(self, filepath: str):
        """فتح ملف Word بالتطبيق الافتراضي"""
        try:
            if os.path.exists(filepath):
                if os.name == 'nt':  # Windows
                    os.startfile(filepath)
                elif os.name == 'posix':  # Linux/Mac
                    subprocess.run(['xdg-open', filepath])
                
                logging.info(f"تم فتح ملف Word: {filepath}")
                
                # إظهار رسالة للمستخدم
                if self.parent:
                    QMessageBox.information(
                        self.parent, 
                        "تم إنشاء الملف", 
                        f"تم إنشاء ملف Word بنجاح وفتحه في التطبيق الافتراضي.\n\nمسار الملف:\n{filepath}"
                    )
            else:
                logging.error(f"ملف Word غير موجود: {filepath}")
                if self.parent:
                    QMessageBox.warning(self.parent, "خطأ", "لم يتم العثور على الملف المُنشأ")
                    
        except Exception as e:
            logging.error(f"خطأ في فتح ملف Word: {e}")
            if self.parent:
                QMessageBox.warning(
                    self.parent, 
                    "تعذر فتح الملف", 
                    f"تم إنشاء الملف بنجاح ولكن تعذر فتحه تلقائياً.\n\nيمكنك العثور على الملف في:\n{filepath}"
                )
    
    def create_teacher_salary_details_word(self, teacher: Dict, salary: Dict, statistics_data=None, salaries_data=None):
        """إنشاء تفاصيل راتب المعلم أو الموظف في ملف Word"""
        try:
            # فحص البيانات الأساسية
            if not teacher:
                teacher = {}
            if not salary:
                salary = {}

            # فحص وتنظيف بيانات salaries_data
            if salaries_data is None:
                salaries_data = []
            elif not isinstance(salaries_data, list):
                salaries_data = []

            # تنظيف البيانات من القيم None
            salaries_data = [s for s in salaries_data if s is not None and isinstance(s, dict)]

            # فحص وتنظيف بيانات statistics_data
            if statistics_data is None:
                statistics_data = {}
            elif not isinstance(statistics_data, dict):
                statistics_data = {}

            # إنشاء مستند Word جديد
            doc = Document()
            
            # إعداد اتجاه النص من اليمين لليسار
            from docx.oxml import OxmlElement
            normal_style = doc.styles['Normal']
            pPr = normal_style.element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            rPr = normal_style.element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rPr.append(rtl)
            
            # إعداد الصفحة
            section = doc.sections[0]
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(11)
            
            # إضافة عنوان رئيسي
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run('تقرير مفصل لرواتب الموظف')
            title_run.font.size = Pt(18)
            title_run.bold = True
            title_run.font.name = 'Arial'
            
            # إضافة تاريخ الطباعة
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_paragraph.add_run(f'تاريخ الطباعة: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            date_run.font.size = Pt(12)
            date_run.font.name = 'Arial'
            
            # إضافة شهر الراتب
            month_paragraph = doc.add_paragraph()
            month_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            month_run = month_paragraph.add_run(f'الشهر: {salary.get("month_year", "غير محدد")}')
            month_run.font.size = Pt(14)
            month_run.bold = True
            month_run.font.name = 'Arial'
            
            doc.add_paragraph()
            
            # بيانات الموظف (عرض جميع الحقول مع قيم مناسبة)
            teacher_data = [
                ('الاسم', teacher.get('name', 'غير محدد')),
                ('المدرسة', teacher.get('school_name', 'غير متوفر') if teacher.get('school_name') and teacher.get('school_name') != 'غير محدد' else 'غير متوفر'),
                ('عدد الحصص', teacher.get('class_hours', 'غير محدد') if teacher.get('class_hours') and str(teacher.get('class_hours', '')).strip() not in ['', '0'] else 'غير محدد'),
                ('رقم الهاتف', teacher.get('phone', 'غير متوفر') if teacher.get('phone') and teacher.get('phone') != 'غير محدد' else 'غير متوفر'),
                ('ملاحظات', teacher.get('notes', 'لا توجد ملاحظات') if teacher.get('notes') and teacher.get('notes') not in ['-', 'غير محدد'] else 'لا توجد ملاحظات')
            ]

            # إنشاء الجدول بالحجم المناسب
            num_rows = len(teacher_data) + 1  # +1 للعناوين
            teacher_table = doc.add_table(rows=num_rows, cols=2)
            teacher_table.style = 'Table Grid'
            teacher_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # عناوين الجدول
            headers = ['البيان', 'القيمة']
            hdr_cells = teacher_table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # التأكد من وجود run قبل الوصول إليه
                if hdr_cells[i].paragraphs[0].runs:
                    run = hdr_cells[i].paragraphs[0].runs[0]
                else:
                    run = hdr_cells[i].paragraphs[0].add_run()
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'

            # إضافة بيانات الموظف
            for i, (label, value) in enumerate(teacher_data, 1):
                row = teacher_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)

                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # التأكد من وجود runs قبل الوصول إليها
                    if cell.paragraphs[0].runs:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Arial'
            
            doc.add_paragraph()
            
            # إضافة قسم الإحصائيات إذا كانت متوفرة
            if statistics_data and any(statistics_data.values()):
                # عنوان الإحصائيات
                stats_title = doc.add_paragraph()
                stats_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                stats_run = stats_title.add_run('إحصائيات الرواتب')
                stats_run.font.size = Pt(16)
                stats_run.bold = True
                stats_run.font.name = 'Arial'
                
                doc.add_paragraph()
                
                # جدول الإحصائيات
                stats_table = doc.add_table(rows=1, cols=4)  # بدء بصف واحد فقط
                stats_table.style = 'Table Grid'
                stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # عناوين الإحصائيات
                stats_headers = ['النوع', 'العدد', 'المبلغ', 'المتوسط']
                hdr_cells = stats_table.rows[0].cells
                for i, header in enumerate(stats_headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # التأكد من وجود run قبل الوصول إليه
                    if hdr_cells[i].paragraphs[0].runs:
                        run = hdr_cells[i].paragraphs[0].runs[0]
                    else:
                        run = hdr_cells[i].paragraphs[0].add_run()
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
                
                # دالة آمنة للحصول على القيم
                def safe_get(value, default=0):
                    """الحصول على قيمة آمنة"""
                    if value is None:
                        return default
                    try:
                        return float(value) if isinstance(value, (int, float, str)) else default
                    except (ValueError, TypeError):
                        return default
                
                # إضافة صفوف البيانات
                stats_rows = [
                    ('إجمالي الرواتب', safe_get(statistics_data.get('total_count', 0)), safe_get(statistics_data.get('total_amount', 0)), safe_get(statistics_data.get('average_salary', 0))),
                    ('رواتب هذا العام', safe_get(statistics_data.get('current_year_count', 0)), safe_get(statistics_data.get('current_year_amount', 0)), 0),
                    ('رواتب هذا الشهر', safe_get(statistics_data.get('current_month_count', 0)), safe_get(statistics_data.get('current_month_amount', 0)), 0),
                    ('أعلى راتب', 1, safe_get(statistics_data.get('highest_salary', 0)), 0)
                ]
                
                for label, count, amount, avg in stats_rows:
                    row = stats_table.add_row()  # إضافة صف جديد
                    row.cells[0].text = label
                    row.cells[1].text = str(int(count))
                    row.cells[2].text = f"{amount:,.0f} دينار" if amount > 0 else "0 دينار"
                    row.cells[3].text = f"{avg:,.0f} دينار" if avg > 0 else "-"
                    
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # التأكد من وجود runs قبل الوصول إليها
                        if cell.paragraphs[0].runs:
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Arial'
                
                doc.add_paragraph()
            
            # جدول تفاصيل الراتب
            salary_table = doc.add_table(rows=3, cols=2)  # تم تقليل الصفوف من 5 إلى 3
            salary_table.style = 'Table Grid'
            salary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # عناوين جدول الراتب
            salary_headers = ['تفصيل الراتب', 'المبلغ']
            hdr_cells = salary_table.rows[0].cells
            for i, header in enumerate(salary_headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # التأكد من وجود run قبل الوصول إليه
                if hdr_cells[i].paragraphs[0].runs:
                    run = hdr_cells[i].paragraphs[0].runs[0]
                else:
                    run = hdr_cells[i].paragraphs[0].add_run()
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
            
            # الحصول على آخر راتب من البيانات
            last_salary_info = "غير متوفر"
            if salaries_data and len(salaries_data) > 0:
                # ترتيب البيانات حسب تاريخ الدفع (الأحدث أولاً)
                sorted_salaries = sorted(salaries_data, key=lambda x: x.get('payment_date', ''), reverse=True)
                last_salary = sorted_salaries[0]
                
                payment_date = last_salary.get('payment_date', '')
                paid_amount = last_salary.get('paid_amount', 0)
                
                if payment_date and paid_amount:
                    try:
                        if isinstance(payment_date, str):
                            formatted_date = datetime.strptime(payment_date, '%Y-%m-%d').strftime('%Y-%m-%d')
                        else:
                            formatted_date = payment_date.strftime('%Y-%m-%d') if hasattr(payment_date, 'strftime') else str(payment_date)
                        
                        last_salary_info = f"{formatted_date} - {paid_amount:,.0f} دينار"
                    except:
                        last_salary_info = f"{payment_date} - {paid_amount:,.0f} دينار"
            
            # بيانات الراتب (بدون البدلات والخصومات)
            salary_data = [
                ('الراتب الأساسي', salary.get('basic_salary', 0)),
                ('آخر راتب', last_salary_info)
            ]
            
            for i, (label, amount) in enumerate(salary_data, 1):
                row = salary_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = f"{amount:,.2f} دينار" if isinstance(amount, (int, float)) else str(amount)
                
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # التأكد من وجود runs قبل الوصول إليها
                    if cell.paragraphs[0].runs:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Arial'
                
                # جعل صف آخر راتب بارزاً
                if label == 'آخر راتب':
                    for cell in row.cells:
                        # التأكد من وجود runs قبل الوصول إليها
                        if cell.paragraphs[0].runs:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                                run.font.size = Pt(12)
            
            doc.add_paragraph()
            
            # إضافة جدول الرواتب التفصيلي إذا كانت البيانات متوفرة
            if salaries_data and len(salaries_data) > 0:
                # عنوان جدول الرواتب
                salaries_title = doc.add_paragraph()
                salaries_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                salaries_run = salaries_title.add_run('سجل الرواتب التفصيلي')
                salaries_run.font.size = Pt(16)
                salaries_run.bold = True
                salaries_run.font.name = 'Arial'
                
                doc.add_paragraph()
                
                # إنشاء جدول الرواتب
                num_cols = 8  # المعرف، تاريخ الدفع، المبلغ المدفوع، الراتب الأساسي، من تاريخ، إلى تاريخ، عدد الأيام، الملاحظات
                salaries_table = doc.add_table(rows=1, cols=num_cols)
                salaries_table.style = 'Table Grid'
                salaries_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # عناوين جدول الرواتب
                salaries_headers = ['المعرف', 'تاريخ الدفع', 'المبلغ المدفوع', 'الراتب الأساسي', 'من تاريخ', 'إلى تاريخ', 'عدد الأيام', 'الملاحظات']
                hdr_cells = salaries_table.rows[0].cells
                for i, header in enumerate(salaries_headers):
                    if i < len(hdr_cells):  # فحص إضافي للتأكد من وجود الخلية
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # التأكد من وجود run قبل الوصول إليه
                        if hdr_cells[i].paragraphs[0].runs:
                            run = hdr_cells[i].paragraphs[0].runs[0]
                        else:
                            run = hdr_cells[i].paragraphs[0].add_run()
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                
                # إضافة بيانات الرواتب مع فحص كل سجل
                for salary_record in salaries_data:
                    if not salary_record or not isinstance(salary_record, dict):
                        continue  # تخطي السجلات غير الصحيحة
                    
                    row = salaries_table.add_row()
                    
                    # تنسيق البيانات مع فحص إضافي
                    def format_date(date_str):
                        if not date_str:
                            return ""
                        try:
                            if isinstance(date_str, str):
                                return datetime.strptime(date_str, '%Y-%m-%d').strftime('%Y-%m-%d')
                            else:
                                return date_str.strftime('%Y-%m-%d') if hasattr(date_str, 'strftime') else str(date_str)
                        except:
                            return str(date_str)
                    
                    def format_amount(amount):
                        if amount is None:
                            return "0 دينار"
                        try:
                            return f"{float(amount):,.0f} دينار"
                        except (ValueError, TypeError):
                            return "0 دينار"
                    
                    row_data = [
                        str(salary_record.get('id', '')),
                        format_date(salary_record.get('payment_date', '')),
                        format_amount(salary_record.get('paid_amount', 0)),
                        format_amount(salary_record.get('base_salary', 0)),
                        format_date(salary_record.get('from_date', '')),
                        format_date(salary_record.get('to_date', '')),
                        str(salary_record.get('days_count', '')),
                        salary_record.get('notes', '') or ""
                    ]
                    
                    # فحص أن عدد البيانات يطابق عدد الأعمدة
                    for i, value in enumerate(row_data):
                        if i < len(row.cells):  # فحص إضافي للتأكد من وجود الخلية
                            row.cells[i].text = value
                            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # التأكد من وجود runs قبل الوصول إليها
                            if row.cells[i].paragraphs[0].runs:
                                for run in row.cells[i].paragraphs[0].runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'Arial'
                
                # ضبط عرض الأعمدة
                for column in salaries_table.columns:
                    for cell in column.cells:
                        cell.width = Inches(0.8)
            
            # إضافة تذييل
            doc.add_paragraph()
            footer_paragraph = doc.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_paragraph.add_run('نظام غريديا لادارة حسابات المدارس الاهلية')
            footer_run.font.size = Pt(10)
            footer_run.italic = True
            footer_run.font.name = 'Arial'
            
            # حفظ الملف
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"تقرير_رواتب_{teacher.get('name', 'موظف')}_{timestamp}.docx"
            filepath = os.path.join(temp_dir, filename)
            
            doc.save(filepath)
            
            # فتح الملف
            self.open_word_file(filepath)
            
            logging.info(f"تم إنشاء ملف Word لتفاصيل راتب الموظف: {filepath}")
            return filepath
            
        except Exception as e:
            logging.error(f"خطأ في إنشاء ملف Word لتفاصيل راتب الموظف: {e}")
            logging.error(f"نوع الخطأ: {type(e).__name__}")
            import traceback
            logging.error(f"تتبع الخطأ:\n{traceback.format_exc()}")
            if self.parent:
                QMessageBox.critical(self.parent, "خطأ", f"فشل في إنشاء ملف Word:\n{str(e)}")
            return None


def create_teacher_salary_details_word_document(teacher: Dict, salary: Dict, parent=None, statistics_data=None, salaries_data=None):
    """دالة مساعدة لإنشاء مستند Word لتفاصيل راتب المعلم أو الموظف مع الإحصائيات وجدول الرواتب"""
    try:
        if not DOCX_AVAILABLE:
            if parent:
                QMessageBox.critical(
                    parent, 
                    "مكتبة مفقودة", 
                    "مكتبة python-docx غير مثبتة.\nيرجى تثبيتها لإنشاء ملفات Word."
                )
            return None
        
        word_manager = WordManager(parent)
        return word_manager.create_teacher_salary_details_word(teacher, salary, statistics_data, salaries_data)
        
    except Exception as e:
        logging.error(f"خطأ في إنشاء مستند Word لتفاصيل راتب الموظف: {e}")
        if parent:
            QMessageBox.critical(parent, "خطأ", f"فشل في إنشاء مستند Word:\n{str(e)}")
        return None


def create_students_word_document(students: List[Dict], selected_columns: Dict, filter_info: Optional[str] = None, parent=None):
    """دالة مساعدة لإنشاء مستند Word لقائمة الطلاب"""
    try:
        if not DOCX_AVAILABLE:
            if parent:
                QMessageBox.critical(
                    parent, 
                    "مكتبة مفقودة", 
                    "مكتبة python-docx غير مثبتة.\nيرجى تثبيتها لإنشاء ملفات Word."
                )
            return None
        
        word_manager = WordManager(parent)
        return word_manager.create_students_list_word(students, selected_columns, filter_info)
        
    except Exception as e:
        logging.error(f"خطأ في إنشاء مستند Word لقائمة الطلاب: {e}")
        if parent:
            QMessageBox.critical(parent, "خطأ", f"فشل في إنشاء ملف Word:\n{str(e)}")
        return None
